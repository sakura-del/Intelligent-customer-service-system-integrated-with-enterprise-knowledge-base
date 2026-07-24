"""文档解析器。

按扩展名分发到不同的解析实现，统一产出 ParsedDocument。
覆盖 PDF / Word / HTML / 纯文本 / Markdown 等常见来源，
为后续切分与向量化提供结构化输入。
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from app.core.logging import get_logger
from app.schemas.knowledge import ParsedDocument, ParsedPage, SectionInfo

logger = get_logger("app.knowledge.parsers")

# Markdown 与 HTML 标题正则：用于从纯文本中识别章节边界
_MARKDOWN_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def compute_doc_hash(text: str) -> str:
    """计算文档内容哈希，用于去重与版本追踪。

    使用 sha256 取前 16 位作为短哈希，避免元数据过长。
    """
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def parse_pdf(file_path: Path) -> ParsedDocument:
    """解析 PDF：使用 PyMuPDF 提取每页文本与页码。

    PDF 天然按页组织内容，逐页提取便于后续按页切分与回溯。
    """
    import fitz  # PyMuPDF

    pages: list[ParsedPage] = []
    full_text_parts: list[str] = []
    with fitz.open(file_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            # 去除多余空白行，保留可读性
            cleaned_text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
            pages.append(ParsedPage(page_number=page_index, text=cleaned_text))
            full_text_parts.append(cleaned_text)

    full_text = "\n".join(full_text_parts)
    return ParsedDocument(
        source=file_path.name,
        file_type="pdf",
        pages=pages,
        doc_hash=compute_doc_hash(full_text),
    )


def _extract_markdown_sections(text: str) -> list[SectionInfo]:
    """从纯文本中识别 Markdown 标题作为章节边界。"""
    sections: list[SectionInfo] = []
    for match in _MARKDOWN_HEADING_PATTERN.finditer(text):
        level = len(match.group(1))
        title = match.group(2).strip()
        sections.append(SectionInfo(title=title, level=level))
    return sections


def parse_docx(file_path: Path) -> ParsedDocument:
    """解析 Word：使用 python-docx 提取段落与标题样式。

    Word 没有页码概念，但段落样式 Heading 1/2 可作为章节层级，
    多个段落聚合为一个虚拟页以便统一抽象。
    """
    import docx

    document = docx.Document(str(file_path))
    sections: list[SectionInfo] = []
    paragraph_texts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        # 通过样式名识别标题层级，便于后续按章节聚合
        if style_name.startswith("heading"):
            level = 1
            digits = re.findall(r"\d+", style_name)
            if digits:
                level = int(digits[0])
            sections.append(SectionInfo(title=text, level=level))
        paragraph_texts.append(text)

    full_text = "\n".join(paragraph_texts)
    # Word 无页码，将全部内容作为单一虚拟页处理
    pages = [ParsedPage(page_number=1, text=full_text, sections=sections)]
    return ParsedDocument(
        source=file_path.name,
        file_type="docx",
        pages=pages,
        doc_hash=compute_doc_hash(full_text),
    )


def parse_html(file_path: Path) -> ParsedDocument:
    """解析 HTML：使用 BeautifulSoup 清洗标签，保留语义文本。

    网页内容通常包含导航与广告噪声，通过标签白名单提取正文，
    并依据 h1-h6 标签推断章节结构。
    """
    from bs4 import BeautifulSoup

    raw = file_path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")

    # 移除脚本与样式噪声，避免污染正文
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    sections: list[SectionInfo] = []
    for heading_tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
        title = heading_tag.get_text(strip=True)
        if title:
            level = int(heading_tag.name[1])
            sections.append(SectionInfo(title=title, level=level))

    text = soup.get_text(separator="\n")
    cleaned_text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    pages = [ParsedPage(page_number=1, text=cleaned_text, sections=sections)]
    return ParsedDocument(
        source=file_path.name,
        file_type="html",
        pages=pages,
        doc_hash=compute_doc_hash(cleaned_text),
    )


def parse_plain_text(file_path: Path) -> ParsedDocument:
    """解析纯文本/Markdown：直接读取，并尝试识别 Markdown 标题。"""
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    sections: list[SectionInfo] = (
        _extract_markdown_sections(text) if file_path.suffix.lower() in {".md", ".markdown"} else []
    )
    pages = [ParsedPage(page_number=1, text=text, sections=sections)]
    file_type = "md" if file_path.suffix.lower() in {".md", ".markdown"} else "txt"
    return ParsedDocument(
        source=file_path.name,
        file_type=file_type,
        pages=pages,
        doc_hash=compute_doc_hash(text),
    )


# 扩展名 → 解析函数映射，集中配置便于扩展
_PARSER_REGISTRY = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".html": parse_html,
    ".htm": parse_html,
    ".txt": parse_plain_text,
    ".md": parse_plain_text,
    ".markdown": parse_plain_text,
}


def parse_file(file_path: str | Path) -> ParsedDocument:
    """按扩展名分发解析入口。

    工厂模式：调用方无需关心具体格式，统一拿到 ParsedDocument。
    未知扩展名回退为纯文本解析，保证流水线不中断。
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"待解析文件不存在：{path}")

    suffix = path.suffix.lower()
    parser = _PARSER_REGISTRY.get(suffix, parse_plain_text)
    logger.info("解析文件 %s（类型 %s）使用 %s", path.name, suffix, parser.__name__)
    try:
        return parser(path)
    except Exception as exc:
        # 解析失败时返回空文档并向上抛错，由流水线决定是否中断
        logger.exception("解析文件 %s 失败：%s", path.name, exc)
        raise
